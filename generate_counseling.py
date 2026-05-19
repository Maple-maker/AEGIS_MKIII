"""
Generate SSG Moore Initial Counseling — AR 25-50 Memorandum Format
Output: ~/Desktop/SSG_Moore_Initial_Counseling.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import copy


def set_font(run, bold=False, size=12, underline=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline


def add_paragraph(doc, text="", alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
                  size=12, space_before=0, space_after=0, underline=False,
                  left_indent=0, first_line_indent=0):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.left_indent = Inches(left_indent)
    if first_line_indent:
        pf.first_line_indent = Inches(first_line_indent)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, size=size, underline=underline)
    return p


def add_run(para, text, bold=False, size=12, underline=False):
    run = para.add_run(text)
    set_font(run, bold=bold, size=size, underline=underline)
    return run


def set_doc_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_section_header(doc, number, title, space_before=6):
    p = add_paragraph(doc, space_before=space_before, space_after=2)
    add_run(p, f"{number}.  ", bold=True)
    add_run(p, title, bold=True, underline=True)
    add_run(p, ".", bold=True)
    return p


def add_sub(doc, letter, text, indent=0.5):
    p = add_paragraph(doc, space_before=2, space_after=2, left_indent=indent)
    add_run(p, f"{letter}.  ")
    add_run(p, text)
    return p


def build_counseling():
    doc = Document()
    set_doc_margins(doc)

    # Remove default paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(12)

    # ── LETTERHEAD ─────────────────────────────────────────────────────────────
    p = add_paragraph(doc, "DEPARTMENT OF THE ARMY", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
    add_paragraph(doc, "2nd Battalion, 55th Air Defense Artillery Regiment",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "Bravo Battery",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "Camp Mudaysis, Jordan",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc)

    # ── OFFICE SYMBOL / DATE ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    add_run(p, "AFVL-DRB")
    # Tab stop to right-align date — use a table for reliability
    # Use two-cell table (no borders) for office symbol + date alignment
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    # Remove borders
    for cell in tbl.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    left_cell = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]

    lp = left_cell.paragraphs[0]
    lp.clear()
    add_run(lp, "AFVL-DRB")

    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(rp, "16 May 2026")

    add_paragraph(doc)

    # ── MEMORANDUM FOR / SUBJECT ───────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "MEMORANDUM FOR RECORD", bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "SUBJECT: ", bold=True)
    add_run(p, "Initial Counseling of Staff Sergeant [FIRST] [MI]. Moore")

    add_paragraph(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — PURPOSE
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(doc, "1", "Purpose")

    add_sub(doc, "a",
        "The purpose of this counseling is to establish expectations, standards, and "
        "responsibilities for Staff Sergeant [FIRST] [MI]. Moore (hereinafter \"SSG Moore\") "
        "as a Section Chief within Bravo Battery, 2-55 ADAR. This counseling serves as the "
        "foundation of our professional relationship and sets the conditions for mission "
        "success throughout the deployment and beyond.")

    add_sub(doc, "b",
        "This counseling is conducted in accordance with AR 623-3, FM 6-22, and the standards "
        "of the Army Profession. It is not disciplinary in nature. It is a leadership tool "
        "used to align effort, build trust, and develop Soldiers.")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — ROLE AND RESPONSIBILITIES
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(doc, "2", "Role of the Staff Sergeant")

    add_sub(doc, "a",
        "SSG Moore serves as a Section Chief and is the primary link between the platoon "
        "leadership and the Soldiers in the section. This role carries significant "
        "responsibility — both in technical expertise and in the development and welfare "
        "of subordinates.")

    add_sub(doc, "b",
        "SSG Moore is expected to embody the BE/KNOW/DO framework in all professional "
        "actions. CHARACTER (BE) must come first. Competence (KNOW) and presence (DO) "
        "follow naturally from a strong character foundation. A leader who performs the "
        "right actions for the wrong reasons is not yet the leader the Army needs.")

    add_sub(doc, "c",
        "SSG Moore is the first line of care, discipline, and accountability for every "
        "Soldier in the section. Any issue that affects a Soldier's readiness, welfare, "
        "or performance is SSG Moore's issue — not just something to pass up the chain.")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — EXPECTATIONS AND STANDARDS
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(doc, "3", "Expectations and Standards")

    add_sub(doc, "a",
        "ACCOUNTABILITY. Formation accountability is non-negotiable. SSG Moore is responsible "
        "for knowing the location, status, and disposition of every assigned Soldier at all "
        "times. Accountability reports will be accurate, on time, and passed through the chain "
        "of command without prompting.")

    add_sub(doc, "b",
        "STANDARDS ENFORCEMENT. SSG Moore will enforce all Army standards — appearance, "
        "conduct, physical fitness, and performance — firmly and consistently. Inconsistent "
        "standards enforcement is a leadership failure. The standard is not what SSG Moore "
        "tolerates; the standard is what the Army requires.")

    add_sub(doc, "c",
        "COUNSELING. SSG Moore will conduct timely, documented counselings for all assigned "
        "Soldiers IAW FM 6-22. Initial counselings will be completed within the first 30 days "
        "of assignment or rating period. Event-driven counselings will be conducted within "
        "72 hours of the event. Monthly performance counselings will be conducted and "
        "documented on DA Form 4856.")

    add_sub(doc, "d",
        "ADMINISTRATIVE REQUIREMENTS. SSG Moore will ensure all administrative actions — "
        "5988-Es, DA 6, training schedules, MEDPROS updates, pass requests, and leave "
        "documents — are completed accurately and on time. Errors and omissions in "
        "administrative work reflect directly on section leadership.")

    add_sub(doc, "e",
        "COMMUNICATION. SSG Moore will maintain two-way communication with both subordinates "
        "and leadership. Issues will be reported up the chain of command early — not when "
        "they become crises. SSG Moore is expected to anticipate problems, not just react "
        "to them.")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — SAFETY
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(doc, "4", "Safety")

    add_sub(doc, "a",
        "In a deployed, threat-environment context, force protection is a leadership "
        "responsibility — not a checklist. SSG Moore will ensure all section personnel "
        "understand and comply with all force protection measures, emergency procedures, "
        "and battle drills relevant to the operational environment.")

    add_sub(doc, "b",
        "Composite risk management will be integrated into all section-level operations "
        "and training. SSG Moore will conduct risk assessments for all non-routine activities "
        "and brief Soldiers accordingly. A Soldier who is not briefed on the risk cannot "
        "mitigate it.")

    add_sub(doc, "c",
        "Fratricide prevention, ROE adherence, and situational awareness are standing "
        "requirements. SSG Moore will ensure all assigned Soldiers maintain proficiency "
        "in engagement procedures, ROE, and emergency action drills.")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — PERFORMANCE AND ACCOUNTABILITY
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(doc, "5", "Performance and Accountability")

    add_sub(doc, "a",
        "SSG Moore will be evaluated on the performance of the section, not merely personal "
        "performance. A section chief who excels individually but whose Soldiers fail to "
        "meet standards has not met the standard of the position.")

    add_sub(doc, "b",
        "NCOER input will be submitted in a timely manner, accurately reflecting observed "
        "performance. SSG Moore will not wait until the rating period closes to begin "
        "documenting Soldier performance. The counseling record is the NCOER.")

    add_sub(doc, "c",
        "Physical readiness is a combat multiplier. SSG Moore will maintain personal ACFT "
        "standards and will develop a section culture in which physical readiness is treated "
        "as a professional obligation, not a personal preference.")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — PROFESSIONAL GROWTH AND DEVELOPMENT
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(doc, "6", "Professional Growth and Development")

    add_sub(doc, "a",
        "The Army's three-domain model for leader development — operational, institutional, "
        "and self-development — applies at every grade. SSG Moore is expected to pursue "
        "growth in all three domains throughout this assignment.")

    add_sub(doc, "b",
        "SSG Moore should identify and actively pursue the next institutional requirement "
        "for career progression: Advanced Leaders Course (ALC) completion, SSD completion, "
        "and any functional courses relevant to the AD specialty. I will support these "
        "goals and help remove barriers where possible.")

    add_sub(doc, "c",
        "Self-development is a personal responsibility. SSG Moore is encouraged to read "
        "professionally, seek mentorship from senior NCOs and officers, and engage "
        "constructively with Army doctrine. Recommended reading: ADP 6-22, FM 6-22, "
        "and the NCO Common Core Competencies.")

    add_sub(doc, "d",
        "I will conduct monthly professional development conversations with SSG Moore, "
        "separate from event-driven counselings. These conversations are not evaluations — "
        "they are investments in the leader SSG Moore is becoming.")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — ARMY VALUES AND PROFESSIONAL CONDUCT
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(doc, "7", "Army Values and Professional Conduct")

    add_sub(doc, "a",
        "LDRSHIP is not a poster on the wall — it is the professional identity of every "
        "Soldier in the United States Army. SSG Moore is expected to model all seven values "
        "in actions, decisions, and relationships. The values are not aspirational; "
        "they are the floor.")

    add_sub(doc, "b",
        "INTEGRITY is non-negotiable. SSG Moore will be honest with subordinates, peers, "
        "and superiors — even when it is uncomfortable. A leader who tells leadership what "
        "they want to hear rather than what is true has failed the mission before it begins.")

    add_sub(doc, "c",
        "LOYALTY flows in all directions: to the mission, to subordinates, to leadership, "
        "and to the institution. SSG Moore will defend subordinates from unjust treatment "
        "while simultaneously holding them to the standard. This is not a contradiction; "
        "it is the definition of the NCO role.")

    add_sub(doc, "d",
        "The Army has zero tolerance for sexual harassment, sexual assault, hazing, "
        "bullying, and retaliation against reporters. SSG Moore will enforce these standards "
        "without exception and will immediately report any violations up the chain of command.")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 8 — ENDURING GUIDANCE
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(doc, "8", "Enduring Guidance")

    add_sub(doc, "a",
        "MOVE TO FRICTION. When things get hard — administratively, operationally, "
        "interpersonally — do not avoid the problem. Move toward it. Leaders who wait for "
        "friction to resolve itself create larger problems downstream. Address issues early, "
        "directly, and with the intent to fix rather than assign blame.")

    add_sub(doc, "b",
        "PROTECT YOUR PEOPLE. The welfare of your Soldiers is your primary responsibility "
        "outside of mission accomplishment. Know what is going on in their lives. Identify "
        "Soldiers who are struggling before they reach a crisis. A section chief who only "
        "knows the tactical picture and not the human picture is only half-present.")

    add_sub(doc, "c",
        "EARN THE STANDARD EVERY DAY. Rank is a responsibility, not a privilege. The right "
        "to lead is earned through consistent demonstration of competence, character, and "
        "commitment. SSG Moore was promoted because of potential — what happens next is "
        "determined by performance.")

    add_sub(doc, "d",
        "MY DOOR IS OPEN. I expect SSG Moore to bring problems to me when they arise — "
        "not after they compound. I will not mistake early reporting for weakness. I will "
        "mistake late reporting for a failure to communicate. We succeed or fail together "
        "as a team.")

    # ══════════════════════════════════════════════════════════════════════════
    # CLOSING
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(doc, "9", "Closing Statement")

    add_sub(doc, "a",
        "This counseling is the beginning of our professional relationship. My intent is "
        "to develop SSG Moore into a more capable leader, a stronger NCO, and a more "
        "effective section chief. I am invested in SSG Moore's success. The standards "
        "outlined above are not bureaucratic requirements — they are the conditions "
        "necessary to take care of Soldiers and accomplish the mission.")

    add_sub(doc, "b",
        "I am available to discuss any aspect of this counseling. I expect SSG Moore to "
        "ask questions, raise concerns, and engage as a professional. A counseling that "
        "ends with no questions is a counseling that was not taken seriously.")

    add_paragraph(doc)

    # ── ACKNOWLEDGEMENT LINE ───────────────────────────────────────────────────
    p = add_paragraph(doc, space_before=4)
    add_run(p, "SSG Moore acknowledges receipt of this counseling and understands its contents.")

    add_paragraph(doc)
    add_paragraph(doc)

    # ── SIGNATURE BLOCK (3-column table) ──────────────────────────────────────
    sig_tbl = doc.add_table(rows=6, cols=3)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_tbl.style = "Table Grid"

    # Remove all borders
    for row in sig_tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)

    def sig_cell(tbl, row, col, text, bold=False, underline=False,
                 align=WD_ALIGN_PARAGRAPH.LEFT):
        cell = tbl.rows[row].cells[col]
        p = cell.paragraphs[0]
        p.clear()
        p.alignment = align
        run = p.add_run(text)
        set_font(run, bold=bold, underline=underline)

    # Row 0: labels
    sig_cell(sig_tbl, 0, 0, "RATED NCO:", bold=True)
    sig_cell(sig_tbl, 0, 1, "SENIOR RATER:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    sig_cell(sig_tbl, 0, 2, "RATER:", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Row 1: blank (signature space)
    for col in range(3):
        sig_cell(sig_tbl, 1, col, "")

    # Row 2: blank (signature space)
    for col in range(3):
        sig_cell(sig_tbl, 2, col, "")

    # Row 3: names
    sig_cell(sig_tbl, 3, 0, "[FIRST] [MI]. MOORE", underline=True)
    sig_cell(sig_tbl, 3, 1, "JAIDEN D. RABATIN", underline=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    sig_cell(sig_tbl, 3, 2, "[FIRST] [MI]. RIVERA", underline=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Row 4: grade / title
    sig_cell(sig_tbl, 4, 0, "SSG, USA")
    sig_cell(sig_tbl, 4, 1, "1LT, AD", align=WD_ALIGN_PARAGRAPH.CENTER)
    sig_cell(sig_tbl, 4, 2, "SFC, USA", align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Row 5: position
    sig_cell(sig_tbl, 5, 0, "Section Chief")
    sig_cell(sig_tbl, 5, 1, "LCHR PLT LDR", align=WD_ALIGN_PARAGRAPH.CENTER)
    sig_cell(sig_tbl, 5, 2, "Platoon Sergeant", align=WD_ALIGN_PARAGRAPH.RIGHT)

    add_paragraph(doc)

    # ── DATE LINES ─────────────────────────────────────────────────────────────
    date_tbl = doc.add_table(rows=1, cols=3)
    date_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    date_tbl.style = "Table Grid"
    for cell in date_tbl.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def date_cell(tbl, col, text, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell = tbl.rows[0].cells[col]
        p = cell.paragraphs[0]
        p.clear()
        p.alignment = align
        run = p.add_run(text)
        set_font(run)

    date_cell(date_tbl, 0, "Date: ______________")
    date_cell(date_tbl, 1, "Date: ______________", align=WD_ALIGN_PARAGRAPH.CENTER)
    date_cell(date_tbl, 2, "Date: ______________", align=WD_ALIGN_PARAGRAPH.RIGHT)

    add_paragraph(doc)

    # ── POC LINE ───────────────────────────────────────────────────────────────
    p = add_paragraph(doc, space_before=4)
    add_run(p, "POC is 1LT Jaiden D. Rabatin, LCHR PLT LDR, at DSN: [NUMBER] or rabatinj6@gmail.com.")

    # ── SIGNATURE ──────────────────────────────────────────────────────────────
    add_paragraph(doc)
    add_paragraph(doc)
    p = add_paragraph(doc)
    add_run(p, "JAIDEN D. RABATIN", underline=True)
    p = add_paragraph(doc)
    add_run(p, "1LT, AD")
    p = add_paragraph(doc)
    add_run(p, "LCHR PLT LDR")

    # ── SAVE ────────────────────────────────────────────────────────────────────
    out = Path.home() / "Desktop" / "SSG_Moore_Initial_Counseling.docx"
    doc.save(str(out))
    print(f"Saved: {out}")


if __name__ == "__main__":
    build_counseling()
